from pathlib import Path
from typing import Union

def read_text_from_any(path: Union[str, Path]) -> str:
    path = Path(path)
    suffix = path.suffix.lower()
    if suffix == '.txt':
        return path.read_text(encoding='utf-8', errors='ignore')
    elif suffix == '.docx':
        from docx import Document
        doc = Document(str(path))
        return '\n\n'.join(p.text for p in doc.paragraphs)
    elif suffix == '.pdf':
        from pdfminer.high_level import extract_text
        return extract_text(str(path)) or ''
    elif suffix == '.rtf':
        from striprtf.striprtf import rtf_to_text
        data = path.read_text(encoding='utf-8', errors='ignore')
        return rtf_to_text(data) or ''
    elif suffix == '.doc':
        raise RuntimeError('.doc should be converted to .docx before reading')
    else:
        raise RuntimeError(f'Unsupported file type: {suffix}')

def write_text_to_docx(text: str, out_path: Union[str, Path]) -> None:
    from docx import Document
    doc = Document()
    for para in text.split('\n\n'):
        doc.add_paragraph(para)
    doc.save(str(out_path))
